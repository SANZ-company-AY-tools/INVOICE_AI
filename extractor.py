"""
Invoice data extraction module using Claude AI vision.
Supports any invoice format - Spanish, international, handwritten, etc.
"""

import os
import json
import base64
import anthropic
from typing import Dict, List, Optional
from pdf2image import convert_from_path
from PIL import Image
from io import BytesIO
from dotenv import load_dotenv
from docx import Document
from docx.shared import Inches
import tempfile

# Load environment variables
load_dotenv()


class InvoiceExtractor:
    """Extract structured data from invoice images and PDFs using Claude AI."""

    def __init__(self):
        api_key = os.getenv('ANTHROPIC_API_KEY')
        if not api_key:
            raise ValueError(
                "ANTHROPIC_API_KEY not found. "
                "Please set it in your .env file or environment variables."
            )
        self.client = anthropic.Anthropic(api_key=api_key)
        self.model = "claude-sonnet-4-20250514"

        # System prompt for invoice extraction
        self.system_prompt = """Eres un experto en extracción de datos de facturas y contabilidad española.
Tu tarea es analizar la imagen de una factura y extraer los siguientes datos en formato JSON:

- company_name: Nombre de la empresa EMISORA (quien emite/vende, NO el cliente)
- tax_id: NIF/CIF del EMISOR (NO del cliente/receptor)
- invoice_number: Número de factura
- date: Fecha de emisión (formato YYYY-MM-DD)
- concept: Concepto breve de la factura
- period_start: Fecha inicio período si aplica (YYYY-MM-DD, null si no)
- period_end: Fecha fin período si aplica (YYYY-MM-DD, null si no)
- currency: Código ISO de la divisa (EUR, USD, BRL, GBP, MXN, etc.)
- base_amount: Base imponible (número decimal)
- tax_rate: Porcentaje IVA/impuesto (número, ej: 21)
- tax_amount: Importe IVA/impuesto (número decimal)
- total: Total factura (número decimal)
- accounting_account: Cuenta contable sugerida del PGC español (solo número, ej: "629")
- accounting_description: Descripción breve de la cuenta (ej: "Otros servicios")

CUENTAS CONTABLES COMUNES (PGC España):
- 621: Arrendamientos y cánones (alquileres)
- 622: Reparaciones y conservación
- 623: Servicios profesionales independientes (abogados, consultores, asesores)
- 624: Transportes
- 625: Primas de seguros
- 626: Servicios bancarios
- 627: Publicidad y propaganda
- 628: Suministros (luz, agua, gas, teléfono, internet)
- 629: Otros servicios
- 600: Compras de mercaderías
- 602: Compras de otros aprovisionamientos
- 606: Descuentos sobre compras
- 631: Otros tributos
- 640: Sueldos y salarios
- 649: Otros gastos sociales

IMPORTANTE:
- Responde SOLO con el JSON, sin texto adicional ni markdown
- tax_id: SOLO el CIF/NIF del EMISOR (quien factura). Si no puedes distinguirlo del receptor con certeza, pon null
- El emisor suele aparecer arriba con su logo, datos fiscales y "Factura emitida por"
- El receptor/cliente suele aparecer como "Facturar a", "Cliente", "Datos del cliente"
- NUNCA pongas el CIF del cliente/receptor. Si tienes dudas, pon null
- currency: Detecta la divisa por símbolos (€, $, R$, £) o códigos. Usa códigos ISO: EUR, USD, BRL, GBP, MXN, CLP, ARS, etc.
- Sugiere la cuenta contable más apropiada según el concepto
- Si no puedes determinar la cuenta, usa "629" (Otros servicios)
- Importes como números decimales (ej: 1234.56)
- Fechas en formato YYYY-MM-DD"""

    def _image_to_base64(self, image: Image.Image, format: str = "PNG") -> str:
        """Convert PIL Image to base64 string."""
        buffer = BytesIO()
        image.save(buffer, format=format)
        return base64.standard_b64encode(buffer.getvalue()).decode("utf-8")

    def _get_media_type(self, file_path: str) -> str:
        """Get media type from file extension."""
        ext = os.path.splitext(file_path)[1].lower()
        media_types = {
            '.png': 'image/png',
            '.jpg': 'image/jpeg',
            '.jpeg': 'image/jpeg',
            '.gif': 'image/gif',
            '.webp': 'image/webp',
            '.bmp': 'image/png',  # Convert BMP to PNG
            '.tiff': 'image/png',  # Convert TIFF to PNG
        }
        return media_types.get(ext, 'image/png')

    def _load_image_as_base64(self, file_path: str) -> tuple[str, str]:
        """Load image file and return base64 data and media type."""
        image = Image.open(file_path)

        # Convert to RGB if necessary (for formats like RGBA, P, etc.)
        if image.mode in ('RGBA', 'P', 'LA'):
            image = image.convert('RGB')

        # Get appropriate format
        ext = os.path.splitext(file_path)[1].lower()
        if ext in ['.jpg', '.jpeg']:
            format = 'JPEG'
            media_type = 'image/jpeg'
        else:
            format = 'PNG'
            media_type = 'image/png'

        base64_data = self._image_to_base64(image, format)
        return base64_data, media_type

    def _load_pdf_as_base64(self, pdf_path: str) -> List[tuple[str, str]]:
        """Convert PDF pages to base64 images."""
        images = convert_from_path(pdf_path, dpi=200)
        result = []
        for image in images:
            if image.mode != 'RGB':
                image = image.convert('RGB')
            base64_data = self._image_to_base64(image, 'PNG')
            result.append((base64_data, 'image/png'))
        return result

    def _load_docx_as_base64(self, docx_path: str) -> List[tuple[str, str]]:
        """Extract images from Word document or convert to image if no images found."""
        doc = Document(docx_path)
        images_data = []

        # Try to extract embedded images
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                try:
                    image_data = rel.target_part.blob
                    image = Image.open(BytesIO(image_data))
                    if image.mode in ('RGBA', 'P', 'LA'):
                        image = image.convert('RGB')
                    base64_data = self._image_to_base64(image, 'PNG')
                    images_data.append((base64_data, 'image/png'))
                except Exception:
                    continue

        # If no images found, extract text and send as context
        if not images_data:
            # Extract all text from document
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    full_text.append(' | '.join(row_text))

            # Create a simple image with text info for Claude
            # Since we have text, we'll pass it differently
            text_content = '\n'.join(full_text)
            return [('TEXT_CONTENT', text_content)]

        return images_data

    def _call_claude_vision(self, images_data: List[tuple[str, str]]) -> Dict:
        """Call Claude API with images and extract invoice data."""
        # Build content with all images
        content = []
        has_text_content = False
        text_content = ""

        for base64_data, media_type in images_data:
            if base64_data == 'TEXT_CONTENT':
                # This is text from a Word document without images
                has_text_content = True
                text_content = media_type
            else:
                content.append({
                    "type": "image",
                    "source": {
                        "type": "base64",
                        "media_type": media_type,
                        "data": base64_data,
                    }
                })

        if has_text_content:
            content.append({
                "type": "text",
                "text": f"Extrae los datos de esta factura (contenido del documento Word):\n\n{text_content}\n\nDevuelve SOLO el JSON con los campos especificados."
            })
        else:
            content.append({
                "type": "text",
                "text": "Extrae los datos de esta factura y devuelve SOLO el JSON con los campos especificados."
            })

        # Call Claude API
        message = self.client.messages.create(
            model=self.model,
            max_tokens=1024,
            system=self.system_prompt,
            messages=[
                {"role": "user", "content": content}
            ]
        )

        # Parse response
        response_text = message.content[0].text.strip()

        # Clean up response if it contains markdown code blocks
        if response_text.startswith("```"):
            lines = response_text.split("\n")
            # Remove first line (```json or ```) and last line (```)
            response_text = "\n".join(lines[1:-1])

        try:
            return json.loads(response_text)
        except json.JSONDecodeError as e:
            return {
                "error": f"Failed to parse Claude response: {e}",
                "raw_response": response_text
            }

    def extract_data(self, file_path: str) -> Dict:
        """Extract all relevant data from an invoice file."""
        ext = os.path.splitext(file_path)[1].lower()

        try:
            # Load file as base64 image(s)
            if ext == '.pdf':
                images_data = self._load_pdf_as_base64(file_path)
            elif ext in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.gif', '.webp']:
                images_data = [self._load_image_as_base64(file_path)]
            elif ext in ['.docx', '.doc']:
                images_data = self._load_docx_as_base64(file_path)
            else:
                raise ValueError(f"Unsupported file format: {ext}")

            # Call Claude to extract data
            extracted = self._call_claude_vision(images_data)

            # Build result with standard fields
            result = {
                'file_name': os.path.basename(file_path),
                'company_name': extracted.get('company_name'),
                'tax_id': extracted.get('tax_id'),
                'invoice_number': extracted.get('invoice_number'),
                'date': extracted.get('date'),
                'concept': extracted.get('concept'),
                'period_start': extracted.get('period_start'),
                'period_end': extracted.get('period_end'),
                'currency': extracted.get('currency', 'EUR'),
                'base_amount': self._parse_number(extracted.get('base_amount')),
                'tax_rate': self._parse_number(extracted.get('tax_rate')),
                'tax_amount': self._parse_number(extracted.get('tax_amount')),
                'total': self._parse_number(extracted.get('total')),
                'accounting_account': extracted.get('accounting_account', '629'),
                'accounting_description': extracted.get('accounting_description', 'Otros servicios'),
                'status': 'success' if 'error' not in extracted else 'error',
            }

            if 'error' in extracted:
                result['error'] = extracted['error']
                result['status'] = 'error'

            return result

        except Exception as e:
            return {
                'file_name': os.path.basename(file_path),
                'status': 'error',
                'error': str(e)
            }

    def _parse_number(self, value) -> Optional[float]:
        """Parse a value to float, handling various formats."""
        if value is None:
            return None
        if isinstance(value, (int, float)):
            return float(value)
        if isinstance(value, str):
            try:
                # Remove currency symbols and spaces
                cleaned = value.replace('€', '').replace('$', '').replace('£', '').strip()
                # Handle European format (1.234,56)
                if ',' in cleaned and '.' in cleaned:
                    if cleaned.rindex(',') > cleaned.rindex('.'):
                        cleaned = cleaned.replace('.', '').replace(',', '.')
                    else:
                        cleaned = cleaned.replace(',', '')
                elif ',' in cleaned:
                    cleaned = cleaned.replace(',', '.')
                return float(cleaned)
            except (ValueError, AttributeError):
                return None
        return None

    def process_multiple_files(self, file_paths: List[str]) -> List[Dict]:
        """Process multiple invoice files."""
        results = []
        for file_path in file_paths:
            data = self.extract_data(file_path)
            results.append(data)
        return results
